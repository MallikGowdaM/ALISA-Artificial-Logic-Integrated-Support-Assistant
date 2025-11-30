import os
import re
from docx import Document
from groq import Groq
from datetime import datetime

# ------------------- CONFIG -------------------
BACKEND = "ollama"   # Change to "groq" if you want to use Groq API

# ------------------- BACKEND SETUP -------------------
if BACKEND == "ollama":
    import ollama
    MODEL_NAME = "gemma:2b"  # you can change to llama3, mistral, etc.

    def generate(prompt):
        response = ollama.chat(model=MODEL_NAME, messages=[{"role": "user", "content": prompt}])
        return response["message"]["content"]

elif BACKEND == "groq":
    GROQ_API_KEY = os.getenv("GROQ_API_KEY", "your-groq-api-key-here")
    client = Groq(api_key=GROQ_API_KEY)
    MODEL_NAME = "llama3-8b-8192"  # or "llama3-70b-8192"

    def generate(prompt):
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content


# ------------------- HELPER FUNCTIONS -------------------
def format_time(hour):
    """Convert float hour to 12-hour format with AM/PM."""
    h = int(hour)
    m = int((hour - h) * 60)
    period = "AM" if h < 12 else "PM"
    if h > 12:
        h -= 12
    if h == 0:
        h = 12
    return f"{h}:{m:02d} {period}"

def clean_for_speech(text):
    """Remove symbols and extra spaces so speech sounds natural."""
    text = re.sub(r"[-•*:]+", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()

NUM_WORDS = {
    "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
    "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
    "eleven": 11, "twelve": 12
}

def parse_number(text):
    text = text.lower().strip()
    # direct number
    if text.isdigit():
        return int(text)
    # e.g. "5 hours"
    if text.split()[0].isdigit():
        return int(text.split()[0])
    # e.g. "five hours"
    for word, num in NUM_WORDS.items():
        if text.startswith(word):
            return num
    return None

def parse_date_easy(text):
    """
    Parse dates like '12th Sep', '5 October', '1 Jan' into YYYY-MM-DD.
    Defaults to current year if year not spoken.
    """
    text = text.lower().replace("st", "").replace("nd", "").replace("rd", "").replace("th", "")
    current_year = datetime.now().year

    try:
        dt = datetime.strptime(f"{text} {current_year}", "%d %b %Y")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        try:
            dt = datetime.strptime(f"{text} {current_year}", "%d %B %Y")
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            return None


def brief_trip_info(itinerary):
    lines = itinerary.split("\n")
    spots = [line for line in lines if any(word in line.lower() for word in ["visit", "explore", "attraction", "spot", "temple", "park", "museum"])]
    brief = "Here are some spots included: " + ", ".join(spots[:5])
    return clean_for_speech(brief)

def brief_study_info(timetable):
    lines = timetable.split("\n")
    timings = [line for line in lines if any(ch.isdigit() for ch in line)]
    brief = "Your study schedule covers these times: " + ", ".join(timings[:5])
    return clean_for_speech(brief)


def append_to_docx(filename, heading, content):
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(content)
    doc.save(filename)


# ------------------- TRIP PLANNING -------------------
def plan_trip(destination, days=3, preferences=None):
    extra = f"Consider these preferences: {preferences}." if preferences and preferences != "none" else ""
    prompt = f"""
    You are a professional travel planner.
    Create a detailed {days}-day trip plan for {destination}.
    Include:
    - 2-3 attractions per day
    - One evening activity
    - Local food suggestions
    - Best transport options
    - Estimated budget per day in INR
    {extra}
    Make it realistic and easy to read.
    """
    return generate(prompt)


# ------------------- STUDY TIMETABLE -------------------
def create_study_timetable(subjects, hours_per_day, difficult_subjects=None, exam_dates=None):
    difficult_str = ", ".join(difficult_subjects) if difficult_subjects else "None"
    exam_dates_str = "\n".join([f"- {sub}: {date}" for sub, date in exam_dates.items()]) if exam_dates else "None"

    prompt = f"""
    You are an expert study planner.
    Create a daily study timetable for these subjects: {', '.join(subjects)}.
    Total hours per day: {hours_per_day}.

    Prioritize difficult subjects: {difficult_str}.
    Exam dates:
    {exam_dates_str}

    - Break sessions into 45-60 mins with breaks.
    - Include meals and rest.
    - Keep it realistic and flexible.
    """
    return generate(prompt)


# ------------------- INTERACTIVE FUNCTIONS -------------------
def plan_trip_interactive(active_lang, speak, takecommand):
    speak("What is your destination?", active_lang)
    destination = takecommand(active_lang)

    speak("How many days is your trip?", active_lang)
    days = takecommand(active_lang)
    try:
        days_val = parse_number(days)
        days = days_val if days_val else 5
    except Exception:
        days = 5

    speak("Do you have any special preferences like adventure, family, vegetarian food, or budget trip?", active_lang)
    preferences = takecommand(active_lang)

    speak(f"Planning a {days}-day trip to {destination}. Please wait while I create your itinerary.", active_lang)
    itinerary = plan_trip(destination, days, preferences)

    # Append to trip_plans.docx
    append_to_docx("trip_plans.docx", f"{days}-Day Trip to {destination}", itinerary)

    # Speak only brief info
    speak(brief_trip_info(itinerary), active_lang)
    speak("For more details, please check the trip_plans.docx file.", active_lang)
    print(itinerary)


def study_timetable_interactive(active_lang, speak, takecommand):
    speak("How many subjects do you have?", active_lang)
    num = takecommand(active_lang)
    num_val = parse_number(num)
    if num_val:
        num = num_val
        subjects = []
        for i in range(num):
            speak(f"Please tell me subject {i + 1}", active_lang)
            sub = takecommand(active_lang)
            subjects.append(sub)

        speak("Do you have any difficult subjects? Say their names or say no.", active_lang)
        difficult_input = takecommand(active_lang)
        difficult_subjects = []
        if difficult_input != "no":
            difficult_subjects = [s.strip() for s in difficult_input.split() if s in subjects]

        speak("Do you want to set exam dates for prioritizing? Say yes or no.", active_lang)
        exam_dates = {}
        answer = takecommand(active_lang)
        if answer == "yes":
            for sub in subjects:
                speak(f"What is the exam date for {sub}? You can say like '12th Sep' or '5 October'.", active_lang)
                date_text = takecommand(active_lang)
                parsed_date = parse_date_easy(date_text)
                if parsed_date:
                    exam_dates[sub] = parsed_date
                else:
                    speak("Sorry, I couldn’t understand the date. Please try again.", active_lang)

        speak("How many total hours can you study per day?", active_lang)
        hrs = takecommand(active_lang)
        hrs_val = parse_number(hrs)
        if hrs_val:
            result = create_study_timetable(subjects, hrs_val, difficult_subjects, exam_dates)

            # Append to study_timetables.docx
            append_to_docx("study_timetables.docx", "Study Timetable", result)

            # Speak only brief info
            speak(brief_study_info(result), active_lang)
            speak("For more details, please check the study_timetables.docx file.", active_lang)
            print(result)
        else:
            speak("Invalid input for hours. Please provide a number.", active_lang)
    else:
        speak("Invalid input for number of subjects. Please provide a number.", active_lang)
